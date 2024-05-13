# Design Pattern : Factory
# Creating Document Types based on type of String
# The Document will contain : Subject, Content | Body

from abc import ABC, abstractmethod
from docx import Document as word
from PyPDF2 import PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import csv
import os

class Document(ABC):
    @abstractmethod
    def create_document(self):
        pass

class PdfDocument(Document):

    def create_document(self, subject : str, content : str, filename : str):
        filepath = os.path.join(os.getcwd(), filename)
        pdf_writer = PdfWriter()

        page = pdf_writer.add_page()

         
        return filepath    
    


class WordDocument(Document):

    def create_document(self, subject : str, content : str, filename : str) -> str:
        filepath = os.path.join(os.getcwd(), filename)
        doc = word()
        doc.add_heading(subject, level=1)
        doc.add_paragraph(content)
        doc.save(filepath)
        return filepath
        



class CSVDocument(Document):

    def create_document(self, filename : str, subject : str, content: list[str]) -> str:
        filepath = os.path.join(os.getcwd(), filename)
        with open(filepath, 'w', newline='') as csv_file:
            csv_writer = csv.writer(csv_file)
            csv_writer.writerow([subject])
            for line in content:
                csv_writer.writerow([line])  
        return filepath

    



class DocumentMaker(Document):
    
    @staticmethod
    def create_document(doc_type, *args, **kwargs):
        doc_types = {
            "pdf" : PdfDocument().create_document,
            "word" : WordDocument().create_document,
            "csv" : CSVDocument().create_document
        }

        document_creator = doc_types.get(doc_type)
        if document_creator: return document_creator(*args, **kwargs)
        else : raise ValueError("Document not supported")



# Type of usage:

DocumentMaker.create_document("word", subject="Design WORD file",filename="Word Factory Design file",content="A design patterns that can be use in code!")
DocumentMaker.create_document("pdf", subject="Design PDF file",filename="Word Factory Design file",content="A design patterns that can be use in code!")
DocumentMaker.create_document("csv", subject="Design CSV file",filename="Factory Design", content="A design pattern that")
